#!/usr/bin/python3
# -*- coding: utf-8 -*-
# @Date    : 2022-05-09
# @Author  : tianyafu

"""
生成数据库设计说明书
docx库参见:https://python-docx.readthedocs.io/en/latest/
样式参见:https://python-docx.readthedocs.io/en/latest/user/styles-using.html
为章节生成编号参见:https://www.osgeo.cn/python-tutorial/docx-addnum.html
"""

import configparser
import os
import sys
import datetime

from docx import Document
from docx.oxml.ns import qn  # 设置字体
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_BUILTIN_STYLE
from docx.shared import Pt, RGBColor

root_path = os.path.abspath(os.path.join(os.path.dirname(__file__), os.pardir, os.pardir))

# 添加自定义class
class_path = root_path + '/dw/classes'
sys.path.append(class_path)

from TableMetaForManual import TableMetaForManual

# 添加自定义模块到系统路径
lib_path = root_path + '/utils'
sys.path.append(lib_path)

from logutil import Logging
import gputil

logger = Logging().get_logger()

# 加载配置文件
config_path = root_path + '/config/prod.conf'
config = configparser.ConfigParser()
config.read(config_path, encoding="utf-8-sig")
gp_section = "greenplum"
gp_host = config.get(gp_section, "host")
gp_user = config.get(gp_section, "user")
gp_passwd = config.get(gp_section, "passwd")
gp_database = config.get(gp_section, "database")
gp_port = config.get(gp_section, "port")

doc_section = 'document'
doc_name = config.get(doc_section, "doc_name")
doc_title = config.get(doc_section, "doc_title")
level_1_title = config.get(doc_section, "level_1_title")

# 使用微软雅黑作为默认字体
default_font = config.get(doc_section, "font_wei_run_ya_hei_name")
# 宋体
font_song_ti_name = config.get(doc_section, "font_song_ti_name")
# 黑体
font_hei_ti_name = config.get(doc_section, "font_hei_ti_name")
# 二号字体
font_er_hao_size = int(config.get(doc_section, "font_er_hao_size"))
# 三号字体
font_san_hao_size = int(config.get(doc_section, "font_san_hao_size"))
# 表的字段(表头)的字体大小
default_table_column_font_size = int(config.get(doc_section, "default_table_column_font_size"))
# 字号 小五
font_xiao_wu_size = int(config.get(doc_section, "font_xiao_wu_size"))
# 表格的样式
table_style = config.get(doc_section, "table_style")


def get_table_meta_data(conn):
    """
    获取所有表的元数据信息
    :conn: 数据库连接
    """
    sql = """
    SELECT
t1.attnum
,t1.table_name
,t1.table_comment
,upper(t1.column_name) column_name
,upper(t1.data_type) data_type
,(
CASE
WHEN t1.attnotnull = 't' then ''
else 'Y'
end
) attnotnull
,(
CASE
WHEN t2.column_default like 'nextval%%' then ''
WHEN t2.column_default like '%%now%%' then 'now'
WHEN t2.column_default is null then ''
else t2.column_default
end 
) AS default_value
,t1.column_comment
from 
(
SELECT
d.attnum
,A.SCHEMANAME AS schema_name
,A.TABLENAME AS table_name
,obj_description(b.oid) table_comment
,D.ATTNAME AS column_name
,REPLACE(REPLACE(REPLACE(FORMAT_TYPE(D.ATTTYPID, D.ATTTYPMOD),'numeric','number'),'character varying','varchar'),'date','date') AS data_type
,d.attnotnull
,E.DESCRIPTION column_comment
FROM PG_TABLES A
INNER JOIN PG_CLASS B
ON A.TABLENAME = B.RELNAME
LEFT JOIN PG_CATALOG.PG_DESCRIPTION E
ON B.OID = E.OBJOID
LEFT JOIN PG_CATALOG.PG_ATTRIBUTE D
ON D.ATTRELID = E.OBJOID
AND D.ATTNUM = E.OBJSUBID
WHERE 
SCHEMANAME = 'public'
AND D.ATTNUM > 0
) t1
left join 
information_schema.columns t2
on t1.schema_name = t2.table_schema
and t1.table_name = t2.table_name
and t1.column_name = t2.column_name
ORDER BY t1.table_name ,t1.attnum
    """
    result = gputil.select_with_conn(logger, sql, conn)
    result_dict = dict()
    for (attnum, table_name, table_comment, column_name, data_type, attnotnull, default_value,
         column_comment) in result:
        table_meta_list = result_dict.get(table_name, [])
        table_meta_instance = TableMetaForManual()
        table_meta_instance.table_name = table_name
        table_meta_instance.tbl_comment = table_comment
        table_meta_instance.column_name = column_name
        table_meta_instance.column_type_name = data_type
        table_meta_instance.column_comment = column_comment
        table_meta_instance.integer_idx = attnum
        table_meta_instance.allow_null = attnotnull
        table_meta_instance.default_value = default_value
        table_meta_list.append(table_meta_instance)
        result_dict[table_name] = table_meta_list
    return result_dict


def get_all_tables(conn):
    """
    从数据库中查询出所有表的表名
    :conn: 数据库连接
    """
    sql = """
SELECT
distinct 
A.TABLENAME AS table_name
FROM PG_TABLES A
INNER JOIN PG_CLASS B
ON A.TABLENAME = B.RELNAME
LEFT JOIN PG_CATALOG.PG_DESCRIPTION E
ON B.OID = E.OBJOID
LEFT JOIN PG_CATALOG.PG_ATTRIBUTE D
ON D.ATTRELID = E.OBJOID
AND D.ATTNUM = E.OBJSUBID
WHERE 
SCHEMANAME = 'public'
    """
    result = gputil.select_with_conn(logger, sql, conn)
    result_list = []
    for i in result:
        result_list.append(i[0])
    return result_list


def get_black_list():
    """
    获取表的黑名单  这部分表都是一些copy表  测试表之类的临时表
    """
    black_List = []
    black_List.append("back_reform_back_copy1")  # 退半件-改革项关联表 copy表
    black_List.append("dev_pc")  # 义乌平台自动化测试用例
    black_List.append("dim_event_hot_spot_dd_i_alias")  # 关键词表下发临时表(该表不要删)
    black_List.append("dim_focus_person_info_copy")  # 重点人员预警中间表
    black_List.append("dim_focus_person_info_copy1")  # 重点人员预警中间表
    black_List.append("dim_grid_basic_list_f")  # 网格信息表************************** 该表不清楚是否在用 没有注释
    black_List.append("dim_work_order_info_dd_i_20211124")  # 总事件详情表(算法用的表)
    black_List.append("dim_work_order_info_dd_i_20211216")  # 总事件详情表(算法用的表)
    black_List.append("dim_work_order_info_dd_i_copy1")  # 总事件详情表(算法用的表)
    black_List.append("dim_work_order_info_dd_i_test2")  # 总事件详情表(算法用的表)
    black_List.append("dim_work_order_info_dd_i_test2_copy")  # 总事件详情表(算法用的表)
    black_List.append("dim_work_order_info_dd_i_tmp")  # 总事件详情表(算法用的表)
    black_List.append("dim_work_order_info_dd_i_tmp_copy")  # 总事件详情表(算法用的表)
    black_List.append("dm_all_warn_clue_dd_i_20211104")  # 事中预警表(算法用的表)
    black_List.append("dm_all_warn_clue_dd_i_20220104")  # 事中预警表(算法用的表)
    black_List.append("dm_all_warn_clue_dd_i_20220120")  # 事中预警表(算法用的表)
    black_List.append("dm_all_warn_clue_dd_i_copy2")  # 事中预警表(算法用的表)
    black_List.append("dm_all_warn_clue_dd_i_tmp")  # 事中预警表(算法用的表)
    black_List.append("dm_bankrupt_enterprises_dd_f_1")  # 破产企业
    black_List.append("dm_case_registration_dd_f_test")  # 案件登记表
    black_List.append("dm_case_registration_dd_f_up")  # 案件登记表
    black_List.append("dm_enterprise_list_dd_f_test")  # 企业和个体工商户信息表
    black_List.append("dm_focus_on_personnel_information_dd_f_1")  # 重点人员表
    black_List.append("dm_focus_on_personnel_information_dd_f_20211108")  # 重点人员表
    black_List.append("dm_focus_on_personnel_information_dd_f_tmp")  # 重点人员表
    black_List.append("dm_focus_on_place_information_dd_f_tmp")  # 重点场所信息表
    black_List.append("dm_focus_person_risk_dd_i_copy")  # 重点人员预警
    black_List.append("dm_focus_person_risk_dd_i_copy1")  # 重点人员预警
    black_List.append("dm_four_platform_call_back_result_dd_i_bak_tianyafu_20210917")  # 四平台预警事件处理情况反馈结果表
    black_List.append("dm_four_platform_push_log_yy_i_bak")  # 四平台推送日志表(记录每一次推送)
    black_List.append("dm_four_platform_push_log_yy_i_bak_20210917_test")  # 四平台推送日志表(记录每一次推送)
    black_List.append("dm_labor_dispute_dd_i_test")  # 劳资纠纷表
    black_List.append("dm_labor_dispute_risk_enterprise_dd_i_1")  # 劳资纠纷风险企业
    black_List.append("dm_labor_dispute_risk_enterprise_dd_i_20210915")  # 劳资纠纷风险企业
    black_List.append("dm_labor_dispute_risk_enterprise_dd_i_20211122")  # 劳资纠纷风险企业
    black_List.append("dm_labor_dispute_risk_enterprise_dd_i_copy")  # 劳资纠纷风险企业
    black_List.append("dm_labor_dispute_risk_enterprise_dd_i_copy1")  # 劳资纠纷风险企业
    black_List.append("dm_labor_dispute_risk_enterprise_dd_i_copy2")  # 劳资纠纷风险企业
    black_List.append("dm_labor_dispute_risk_enterprise_dd_i_tmp")  # 劳资纠纷风险企业
    black_List.append("dm_prepaid_card_dd_iold")  # 预付卡-事件表
    black_List.append("dm_prepaid_card_risk_enterprise_dd_i_2")  # 预付卡-事件表
    black_List.append("dm_prepaid_card_risk_enterprise_dd_i_tmp")  # 预付卡-事件表
    black_List.append("dm_questions_visits_f_20210812")  # 党员两问大走访
    black_List.append("dm_questions_visits_f_last")  # 党员两问大走访
    black_List.append("dm_questions_visits_f_old")  # 党员两问大走访
    black_List.append("dm_questions_visits_f_tmp")  # 党员两问大走访
    black_List.append("dm_safety_index_f_2021_09_09")  # 平安指数
    black_List.append("dm_safety_index_f_20210812")  # 平安指数
    black_List.append("dm_safety_index_f_old")  # 平安指数
    black_List.append("dm_safety_index_f_original_vision_by_liuyang")  # 平安指数
    black_List.append("dm_safety_index_f_test")  # 平安指数
    black_List.append("dm_safety_statistics_20210728")  # 平安指数统计表
    black_List.append("dm_safety_statistics_20210728_before_dx_changed")  # 平安指数统计表
    black_List.append("dm_safety_statistics_test")  # 平安指数统计表
    black_List.append("dm_work_order_info_dd_i_20210621")  # 总事件详情表
    black_List.append("dm_work_order_info_dd_i_20210716")  # 总事件详情表
    black_List.append("dm_work_order_info_dd_i_20220103")  # 总事件详情表
    black_List.append("dm_work_order_info_dd_i_20220104")  # 总事件详情表
    black_List.append("dm_work_order_info_dd_i_alias")  # 总事件详情表
    black_List.append("dm_work_order_info_dd_i_test")  # 总事件详情表
    black_List.append("dm_work_order_info_dd_i_tmp")  # 总事件详情表
    black_List.append("dm_work_order_info_expl_by_hot_spot_dd_i")  # 总事件详情表（根据热点词炸开）
    black_List.append("dm_work_order_info_expl_by_hot_spot_dd_i_20210629")  # 总事件详情表（根据热点词炸开）
    black_List.append("dm_work_order_info_tmp")  # 总事件详情表
    black_List.append("dwd_safety_index_f_four")  # 平安指数
    black_List.append("ods_dazwfw_hj_xzcf_punish_publicity_data_valid_old_jh_yiw_dd_f")  # 不清楚有没有用  但是是ods表
    black_List.append("ods_dsc_dazwfw_hj_xzcf_valid_old_jh_yiw_dd_f")  # 不清楚有没有用  但是是ods表
    black_List.append("ods_focus_on_place_information_dd_i")  # 不清楚有没有用  但是是ods表
    black_List.append("ods_focus_on_place_information_dd_i_test")  # 不清楚有没有用  但是是ods表
    black_List.append("ods_import_person_yy_f_test")  # 重点人员表
    black_List.append("test")  # 重点人员表
    black_List.append("warn_level_info_test")  # 风险等级
    black_List.append("warn_level_info_tmp")  # 风险等级
    black_List.append("yjf20210909")  # 平安指数越级访20210909
    black_List.append("zyx_test")  # 平安指数越级访20210909

    return black_List


def get_tables(all_tables, black_list):
    """
    获取去掉黑名单后的表 即需要体现在数据库设计说明书中的表
    :all_tables: 数据库中所有表的list
    :black_list: 黑名单表的list
    """
    result_list = []
    if len(all_tables) == 0:
        return result_list
    # 将黑名单转成字典
    black_dict = dict()
    for i in black_list:
        black_dict[i] = i
    # 不在黑名单中的表 即需要体现在数据库设计说明书中的表
    for tbl_name in all_tables:
        black_info = black_dict.get(tbl_name, None)
        if black_info is None:
            result_list.append(tbl_name)

    return result_list


def get_document(doc_title, level_1_title):
    """
    获取document对象 设置document的样式 以及设置文档标题和文档一级标题
    """
    doc = Document()
    # "Normal"表示正文的样式，[“Heading 2”]表示2级标题的样式    一级标题的样式关键字为[“Heading 1”]
    # 正文 黑色宋体
    doc.styles["Normal"].font.name = default_font
    doc.styles["Normal"].font.color.rgb = RGBColor(0, 0, 0)

    # 一级标题 黑色宋体
    doc.styles["Heading 1"].font.name = default_font
    doc.styles["Heading 1"].font.color.rgb = RGBColor(0, 0, 0)

    # 二级标题 黑色宋体
    doc.styles["Heading 2"].font.name = default_font
    doc.styles["Heading 2"].font.color.rgb = RGBColor(0, 0, 0)

    # 三级标题 黑色宋体
    doc.styles["Heading 3"].font.name = default_font
    doc.styles["Heading 3"].font.color.rgb = RGBColor(0, 0, 0)

    # 第一大段落 为文档的标题
    paragraph_1 = doc.add_paragraph()
    run_1 = paragraph_1.add_run(doc_title)
    # 设置字体为default_font 大小为二号
    run_1.font.size = Pt(font_er_hao_size)
    run_1.font.name = default_font
    run_1.element.rPr.rFonts.set(qn('w:eastAsia'), default_font)
    # 居中
    paragraph_1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 第二大段落 为文档的一级标题
    paragraph_2 = doc.add_paragraph()
    paragraph_2.style = doc.styles["Heading 1"]
    run_2 = paragraph_2.add_run(level_1_title)
    # 设置字体为宋体 大小为三号
    run_2.font.size = Pt(font_san_hao_size)
    run_2.font.name = font_song_ti_name
    run_2.element.rPr.rFonts.set(qn('w:eastAsia'), font_song_ti_name)

    return doc


def write_tables2(document, table_meta_dict):
    """
    该方法为备份方法 本脚本中不使用
    """
    for table_name in get_tables(all_tables, black_list):
        table_meta_list = table_meta_dict.get(table_name, [])
        if len(table_meta_list) >= 1:
            # 表注释
            tbl_comment = table_meta_list[0].tbl_comment
            title = "%s %s" % (table_name, tbl_comment)
            # 每张表的标题 是二级标题
            # 第二大段落
            paragraph = document.add_paragraph()
            paragraph.style = document.styles["Heading 2"]
            run = paragraph.add_run(title)
            # 设置字体为default_font 大小为小三号
            run.font.size = Pt(15)
            run.font.name = default_font
            run.element.rPr.rFonts.set(qn('w:eastAsia'), default_font)

            # header = document.add_heading(title, 2)
            # 先生成一张表 的表头 为 1行 6列
            table = document.add_table(rows=1, cols=6, style='Table Grid')

            heading_cells = table.rows[0].cells
            heading_cells[0].text = '序号'
            heading_cells[1].text = '字段名称'
            heading_cells[2].text = '字段类型'
            heading_cells[3].text = '允许空'
            heading_cells[4].text = '缺省值'
            heading_cells[5].text = '字段描述'
            for table_meta_for_manual in table_meta_list:
                cells = table.add_row().cells
                cells[0].text = str(table_meta_for_manual.integer_idx)
                cells[1].text = table_meta_for_manual.column_name
                cells[2].text = table_meta_for_manual.column_type_name
                cells[3].text = table_meta_for_manual.allow_null
                cells[4].text = table_meta_for_manual.default_value
                cells[5].text = table_meta_for_manual.column_comment
            for row in table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size = Pt(9)
        else:
            logger.warning("表%s找不到元数据。。。。。。" % table_name)


def set_table_cell(cell, text, font_size, font_name, is_header):
    """
    设置表的单元格的内容 包括 字体样式 是否加粗等
    """
    run = cell.paragraphs[0].add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = font_name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    # 表头字体加粗
    if is_header:
        run.font.bold = True


def write_tables(document, table_meta_dict):
    """
    向document中写入表信息
    """
    for table_name in get_tables(all_tables, black_list):
        table_meta_list = table_meta_dict.get(table_name, [])
        if len(table_meta_list) >= 1:
            # 表注释
            tbl_comment = table_meta_list[0].tbl_comment
            title = "%s %s" % (table_name, tbl_comment)
            # 每张表的标题 是二级标题
            paragraph = document.add_paragraph()
            paragraph.style = document.styles["Heading 2"]
            run = paragraph.add_run(title)
            # 设置字体为宋体 大小为小三号
            run.font.size = Pt(15)
            run.font.name = font_song_ti_name
            run.element.rPr.rFonts.set(qn('w:eastAsia'), font_song_ti_name)

            # 先生成一张表 的表头 为 1行 6列
            table = document.add_table(rows=1, cols=6, style=table_style)

            heading_cells = table.rows[0].cells
            # 这里设置表单元格的值不再使用方法write_tables2的api 本方法中使用的方式还可以设置字体及字号  更灵活
            set_table_cell(heading_cells[0], '序号', default_table_column_font_size, font_song_ti_name,
                           is_header=True)
            set_table_cell(heading_cells[1], '字段名称', default_table_column_font_size, font_song_ti_name,
                           is_header=True)
            set_table_cell(heading_cells[2], '字段类型', default_table_column_font_size, font_song_ti_name,
                           is_header=True)
            set_table_cell(heading_cells[3], '允许空', default_table_column_font_size, font_song_ti_name,
                           is_header=True)
            set_table_cell(heading_cells[4], '缺省值', default_table_column_font_size, font_song_ti_name,
                           is_header=True)
            set_table_cell(heading_cells[5], '字段描述', default_table_column_font_size, font_song_ti_name,
                           is_header=True)
            for table_meta_for_manual in table_meta_list:
                cells = table.add_row().cells
                # 宋体 小五号字体
                set_table_cell(cells[0], str(table_meta_for_manual.integer_idx), font_xiao_wu_size,
                               font_song_ti_name, is_header=False)
                set_table_cell(cells[1], table_meta_for_manual.column_name, font_xiao_wu_size,
                               font_song_ti_name, is_header=False)
                set_table_cell(cells[2], table_meta_for_manual.column_type_name, font_xiao_wu_size,
                               font_song_ti_name, is_header=False)
                set_table_cell(cells[3], table_meta_for_manual.allow_null, font_xiao_wu_size,
                               font_song_ti_name, is_header=False)
                set_table_cell(cells[4], table_meta_for_manual.default_value, font_xiao_wu_size,
                               font_song_ti_name, is_header=False)
                set_table_cell(cells[5], table_meta_for_manual.column_comment, font_xiao_wu_size,
                               font_song_ti_name, is_header=False)
        else:
            logger.warning("表%s找不到元数据。。。。。。" % table_name)


def write_no(doc):
    """
    对document中的标题进行编号
    从类docx.styles.style._NumberingStyle可以得知编号样式尚未实现
    所以这里自己编码进行标题编号
    """
    head1 = 0
    head2 = 0
    head3 = 0
    for para in doc.paragraphs:
        style_name = para.style.name
        if style_name == "Heading 1":
            head1 += 1
            for i in range(len(para.runs)):
                para.runs[i].text = para.runs[i].text.replace(para.text, str(head1) + " " + para.text)
            head2 = 0
            head3 = 0
        if style_name == "Heading 2":
            head2 += 1
            for i in range(len(para.runs)):
                para.runs[i].text = para.runs[i].text.replace(para.text,
                                                              str(head1) + "." + str(head2) + " " + para.text)
        if style_name == "Heading 3":
            head3 += 1
            for i in range(len(para.runs)):
                para.runs[i].text = para.runs[i].text.replace(para.text, str(head1) + "." + str(head2) + "." + str(
                    head3) + " " + para.text)


if __name__ == '__main__':
    try:
        # 获取GP连接
        gp_conn = gputil.maintain_conn(logger, None, gp_host, gp_user, gp_passwd, gp_database,
                                       gp_port)
        output_file_name = "%s@%s.docx" % (doc_name, datetime.datetime.now().strftime('%Y%m%d%H%M%S'))
        output_path = "output/%s" % output_file_name
        # 获取document对象 设置document的样式 以及设置文档标题和文档一级标题
        document = get_document(doc_title, level_1_title)
        # 获取所有表的元数据信息
        table_meta_dict = get_table_meta_data(gp_conn)
        # 从数据库中查询出所有表的表名
        all_tables = get_all_tables(gp_conn)
        # 获取表的黑名单  这部分表都是一些copy表  测试表之类的临时表
        black_list = get_black_list()
        # 获取去掉黑名单后的表 即需要体现在数据库设计说明书中的表
        tables = get_tables(all_tables, black_list)
        if len(tables) > 0:
            # 向document中写入表信息
            write_tables(document, table_meta_dict)
            # 对document中的标题进行编号
            write_no(document)
            document.save(output_path)
            logger.info("%s 已生成......" % doc_name)
        else:
            logger.warning("没有表需要生成%s......" % doc_name)
    except Exception as e:
        logger.error(e)
        raise Exception
    finally:
        # 关闭GP连接
        gputil.close(gp_conn)
